from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "penut-operator-integration-plan.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "667085",
    "light_gray": "F2F4F7",
    "border": "D0D5DD",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"], size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        style_run(first, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].width = Inches(widths[i])
        set_cell_shading(header_cells[i], COLORS["light_gray"])
        set_cell_margins(header_cells[i])
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = header_cells[i].paragraphs[0].add_run(header)
        style_run(run, bold=True, color="0B2545")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(str(value))
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color="D9E2F3")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=COLORS["dark_blue"])
    p.add_run(f" {body}")
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Penut Operator Integration Plan")
    style_run(run, bold=True, color="0B2545", size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run(
        "Flow A proposal for connecting the desktop Operator companion app to Penut"
    )
    style_run(r, color=COLORS["muted"], size=11)

    add_callout(
        doc,
        "Recommendation:",
        "Use Penut's existing approval/auth/audit foundations, and add only two new production tables for browser execution: browser_tasks and browser_task_events. Do not edit existing table schemas for the initial implementation.",
    )

    heading(doc, "1. Executive Summary")
    paragraph(
        doc,
        "Flow A keeps Penut as the system of record for task creation and approval, while the Operator desktop app performs browser-only work locally using the assigned user's authenticated browser profile.",
    )
    paragraph(
        doc,
        "The design intentionally separates approval state from browser execution state. Penut's existing approval tables answer whether a human approved an action. The new browser task tables answer whether the local browser automation has claimed, run, completed, failed, or logged progress for that action.",
    )

    heading(doc, "2. Agreed Product Flow")
    for item in [
        "A Penut user or agent creates a browser-only task.",
        "Penut resolves the assigned org member, such as Emmanuel, and creates an approval request/action for that member.",
        "Emmanuel logs into Operator with Penut auth and sees only tasks assigned to him.",
        "Emmanuel opens the task, edits the prompt if needed, and clicks Approve.",
        "Operator claims the approved task to prevent duplicate execution, runs it locally in Emmanuel's browser, and streams progress events back to Penut.",
        "Penut stores the final result/error and shows task activity in both Operator and Penut surfaces.",
    ]:
        numbered(doc, item)

    heading(doc, "3. Existing Tables Reused Without Schema Edits")
    add_table(
        doc,
        ["Existing table", "How it is used", "Why no schema edit is needed"],
        [
            [
                "approval_requests",
                "Approval container for the requested browser task.",
                "It already models approval visibility, requester context, source, status, and participants.",
            ],
            [
                "approval_actions",
                "Stores the prepared action with opKey such as operator:browser_task:run and params containing the original prompt.",
                "Existing status/result/error fields are enough for approval decision and final execution outcome.",
            ],
            [
                "approval_participants",
                "Stores the assigned member who must approve/run the task.",
                "The assigned user maps naturally to an approval participant.",
            ],
            [
                "audit_logs",
                "Stores important backend/security milestones such as task created, approved, completed, or failed.",
                "Audit events are not the same as user-facing browser activity logs.",
            ],
            [
                "OAuth/auth tables",
                "Operator uses existing device-code/OAuth login and bearer token auth.",
                "Penut already supports this pattern; no new auth tables are needed.",
            ],
        ],
        [1.55, 2.55, 2.4],
    )

    heading(doc, "4. New Table: browser_tasks")
    paragraph(
        doc,
        "browser_tasks is the execution lifecycle record for one browser-only task assigned to one org member.",
    )
    add_table(
        doc,
        ["Field", "Purpose", "Rationale"],
        [
            ["id, uuid", "Internal primary key and public stable identifier.", "Matches existing Penut table conventions."],
            ["org_id, project_id", "Scope task to organization and project.", "Supports tenancy, permissions, and project filtering."],
            ["approval_request_id", "Link to existing approval request.", "Keeps approval context intact."],
            ["approval_action_id", "Link to exact action being approved/executed.", "One browser task should correspond to one approval action."],
            ["requested_by_member_id", "Member who requested/delegated the task.", "Important for audit and delegation history."],
            ["assigned_member_id", "Member whose Operator/browser must execute the task.", "This is how Emmanuel sees tasks assigned to Emmanuel."],
            ["title", "Readable task label.", "Useful for inbox/list UI."],
            ["prompt", "Original natural-language task.", "Keeps task generic and avoids site-specific fields."],
            ["edited_prompt", "User-edited prompt from Operator.", "Preserves original vs final instruction."],
            ["status", "Execution lifecycle.", "Approval state and execution state are different concepts."],
            ["claimed_at, lease_expires_at", "Task claim/lease window.", "Prevents duplicate execution and supports recovery."],
            ["started_at, completed_at", "Run timing.", "Useful for UI and diagnostics."],
            ["result, error", "Structured final outcome.", "Stores final summary/error without overloading logs."],
            ["created_at, updated_at", "Timestamps.", "Standard lifecycle tracking."],
        ],
        [1.65, 2.05, 2.8],
    )

    paragraph(doc, "Recommended status enum:", style="Heading 3")
    for status in [
        "pending_approval",
        "approved",
        "claimed",
        "running",
        "completed",
        "failed",
        "cancelled",
        "rejected",
        "expired",
    ]:
        bullet(doc, status)

    paragraph(doc, "Recommended indexes:", style="Heading 3")
    for idx in [
        "unique(browser_tasks.uuid)",
        "unique(browser_tasks.approval_action_id)",
        "index(assigned_member_id, status, created_at desc)",
        "index(project_id, status, created_at desc)",
        "index(org_id, status, created_at desc)",
        "index(status, lease_expires_at)",
    ]:
        bullet(doc, idx)

    heading(doc, "5. New Table: browser_task_events")
    paragraph(
        doc,
        "browser_task_events stores the user-facing activity timeline for a browser task. This is separate from audit_logs because users need readable execution progress, while audit_logs should remain focused on backend/security/compliance events.",
    )
    add_table(
        doc,
        ["Field", "Purpose", "Rationale"],
        [
            ["id, uuid", "Internal primary key and public stable identifier.", "Matches Penut conventions."],
            ["browser_task_id", "Parent task.", "Models one-to-many task activity correctly."],
            ["org_id, project_id", "Scope event to org/project.", "Supports filtering and authorization checks without repeated joins."],
            ["event_type", "Type such as system, status, browser, agent, result, or error.", "Allows filtering/grouping without parsing message text."],
            ["message", "Friendly activity text.", "UI can show useful logs without raw JSON."],
            ["detail", "Optional structured metadata.", "Keeps timing/result/page data available without polluting message text."],
            ["created_at", "Event timestamp.", "Supports chronological activity display."],
        ],
        [1.65, 2.05, 2.8],
    )
    paragraph(doc, "Recommended indexes:", style="Heading 3")
    for idx in [
        "unique(browser_task_events.uuid)",
        "index(browser_task_id, created_at asc)",
        "index(project_id, created_at desc)",
    ]:
        bullet(doc, idx)

    heading(doc, "6. Tables Explicitly Not Added")
    add_table(
        doc,
        ["Table", "Decision", "Reason"],
        [
            [
                "operator_installations",
                "Do not add for this plan.",
                "Operator login already identifies the member. Device-level management can be added later if multi-device revocation/online status becomes a product requirement.",
            ],
            [
                "operator_runtime_statuses",
                "Do not add for this plan.",
                "Readiness can remain local to Operator for now. Backend runtime visibility can be added later if admin/setup monitoring is needed.",
            ],
        ],
        [1.8, 1.6, 3.1],
    )

    heading(doc, "7. Endpoint Plan")
    add_table(
        doc,
        ["Endpoint", "Purpose", "Notes"],
        [
            ["GET /operator/tasks", "List tasks assigned to the authenticated member.", "Filter by statuses such as pending_approval, running, failed, completed."],
            ["GET /operator/tasks/:id", "Read one assigned task.", "Returns task plus approval/action context needed for review."],
            ["POST /operator/tasks/:id/approve", "Approve and optionally save edited prompt.", "Updates approval_actions.status and browser_tasks.status."],
            ["POST /operator/tasks/:id/reject", "Reject task from Operator.", "Updates approval action and browser task to rejected."],
            ["POST /operator/tasks/:id/claim", "Atomically claim approved task.", "Prevents duplicate execution."],
            ["PATCH /operator/tasks/:id/status", "Update execution status.", "Used for running, completed, failed, cancelled."],
            ["POST /operator/tasks/:id/events", "Append user-facing activity event.", "Stores progress in browser_task_events."],
            ["GET /operator/tasks/:id/events", "Read task activity timeline.", "Used by Operator and optionally Penut web."],
        ],
        [1.9, 2.15, 2.45],
    )

    heading(doc, "8. Example Scenario")
    add_callout(
        doc,
        "Scenario:",
        "Samuel asks Penut: 'Create a task for Emmanuel to send a DM to xyz.' Emmanuel must approve and run it because Emmanuel's browser/session/account will perform the action.",
    )
    for item in [
        "Penut resolves 'Emmanuel' to an active org_members row.",
        "Penut creates approval_requests with Samuel as requester and Emmanuel as participant.",
        "Penut creates approval_actions with opKey operator:browser_task:run and params containing the prompt.",
        "Penut creates browser_tasks with requested_by_member_id = Samuel and assigned_member_id = Emmanuel.",
        "Emmanuel logs into Operator; Operator calls GET /operator/tasks and sees the assigned task.",
        "Emmanuel edits the prompt if needed and clicks Approve.",
        "Backend approves the approval action and moves browser_tasks.status to approved.",
        "Operator claims the task, runs it locally, appends browser_task_events, and updates status to completed or failed.",
        "Penut shows final result plus activity timeline for that task.",
    ]:
        numbered(doc, item)

    heading(doc, "9. Implementation Notes")
    for item in [
        "Do not add site-specific columns such as site, recipient, or target_url. The prompt remains generic and browser-use/AI handles execution planning.",
        "Use existing OAuth device-code flow for Operator login; do not create a parallel auth table.",
        "Use audit_logs only for important backend/security milestones, not every browser step.",
        "Keep approval state and execution state separate: approval_actions.status answers whether the user approved; browser_tasks.status answers what happened during local execution.",
        "Approval should be performed by the assigned member because that member's browser/session executes the task.",
        "Claiming must be atomic to prevent duplicate runs.",
    ]:
        bullet(doc, item)

    heading(doc, "10. Boss Review Decisions Needed")
    for item in [
        "Confirm Flow A as the integration model.",
        "Confirm only two new tables: browser_tasks and browser_task_events.",
        "Confirm no existing table schema edits for the initial implementation.",
        "Confirm browser tasks are assigned to org members via assigned_member_id.",
        "Confirm Operator uses existing OAuth/device-code auth.",
        "Confirm Penut web should display browser task status and activity timeline.",
    ]:
        bullet(doc, item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Penut Operator Integration Plan")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
